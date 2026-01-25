"""
Advanced Multimodal Processing Module
Handles images, documents, audio, video, and real-time media processing
"""

import io
<<<<<<< HEAD
import json
import os
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional, Tuple
=======
from datetime import datetime
>>>>>>> origin/code-quality-refactor-17423438479402428749


class MediaType(Enum):
    """Supported media types"""

    IMAGE = "image"
    DOCUMENT = "document"
    AUDIO = "audio"
    VIDEO = "video"
    TEXT = "text"


class ImageFormat(Enum):
    """Supported image formats"""

    JPEG = "jpeg"
    PNG = "png"
    GIF = "gif"
    WEBP = "webp"
    BMP = "bmp"


class DocumentFormat(Enum):
    """Supported document formats"""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    CSV = "csv"
    JSON = "json"


class AudioFormat(Enum):
    """Supported audio formats"""

    MP3 = "mp3"
    WAV = "wav"
    M4A = "m4a"
    OGG = "ogg"
    FLAC = "flac"


class VideoFormat(Enum):
    """Supported video formats"""

    MP4 = "mp4"
    AVI = "avi"
    MOV = "mov"
    MKV = "mkv"
    WEBM = "webm"


@dataclass
class MediaFile:
    """Represents a processed media file"""

    filename: str
    media_type: MediaType
    format_type: str
    file_size: int
    processed_at: str
    content: Optional[Any] = None
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if not self.processed_at:
            self.processed_at = datetime.now().isoformat()

    def to_dict(self) -> Dict:
        """Convert to dictionary"""
        return {
            "filename": self.filename,
            "media_type": self.media_type.value,
            "format_type": self.format_type,
            "file_size": self.file_size,
            "processed_at": self.processed_at,
            "metadata": self.metadata,
        }


@dataclass
class ImageData:
    """Image-specific data"""

    width: int
    height: int
    channels: int
    color_space: str
    has_alpha: bool
    exif_data: Optional[Dict] = None


@dataclass
class AudioData:
    """Audio-specific data"""

    duration_seconds: float
    sample_rate: int
    channels: int
    bitrate: int
    codec: str
    waveform: Optional[Any] = None


@dataclass
class VideoData:
    """Video-specific data"""

    duration_seconds: float
    fps: int
    resolution: Tuple[int, int]
    codec: str
    bitrate: int
    frame_count: int


class ImageProcessor:
    """Advanced image processing"""

    @staticmethod
    def validate_image(file_path: str) -> Tuple[bool, str]:
        """Validate image file"""
        try:
            from PIL import Image

            img = Image.open(file_path)
            img.verify()
            return True, "Valid image"
        except Exception as e:
            return False, f"Invalid image: {str(e)}"

    @staticmethod
    def extract_image_metadata(file_path: str) -> Dict[str, Any]:
        """Extract image metadata"""
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS

            img = Image.open(file_path)

            metadata = {
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode,
                "channels": len(img.getbands()),
            }

            # Try to extract EXIF data
            try:
                exif_data = img._getexif()
                if exif_data:
                    exif_dict = {}
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_dict[tag] = str(value)
                    metadata["exif"] = exif_dict
            except:
                pass

            return metadata
        except Exception as e:
            return {"error": str(e)}

    @staticmethod
    def resize_image(
        file_path: str, max_width: int = 1920, max_height: int = 1080
    ) -> bytes:
        """Resize image for processing"""
        try:
            from PIL import Image

            img = Image.open(file_path)

            # Calculate scaling
            ratio = min(max_width / img.width, max_height / img.height, 1.0)
            new_width = int(img.width * ratio)
            new_height = int(img.height * ratio)

            # Resize
            img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

            # Convert to bytes
            buffer = io.BytesIO()
            img_resized.save(buffer, format=img.format or "PNG")
            return buffer.getvalue()
        except Exception as e:
            raise Exception(f"Failed to resize image: {str(e)}")

    @staticmethod
    def optimize_for_ai(file_path: str) -> bytes:
        """Optimize image for AI processing"""
        from PIL import Image, ImageEnhance

        try:
            img = Image.open(file_path)

            # Convert to RGB if needed
            if img.mode != "RGB":
                img = img.convert("RGB")

            # Enhance contrast and brightness
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.2)

            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(1.0)

            # Resize if too large
            if img.width > 2048 or img.height > 2048:
                img.thumbnail((2048, 2048), Image.Resampling.LANCZOS)

            buffer = io.BytesIO()
            img.save(buffer, format="PNG", optimize=True)
            return buffer.getvalue()
        except Exception as e:
            raise Exception(f"Failed to optimize image: {str(e)}")


class DocumentProcessor:
    """Advanced document processing"""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF with metadata"""
        try:
            import PyPDF2
            
            text_parts = []
            metadata = {
                "pages": 0,
                "title": "",
                "author": "",
                "created": ""
            }
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata["pages"] = len(reader.pages)

                # Extract metadata
                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title", "")
                    metadata["author"] = reader.metadata.get("/Author", "")
                    metadata["created"] = reader.metadata.get("/CreationDate", "")

                # Extract text
                for page_num, page in enumerate(reader.pages):
                    text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                    text_parts.append(page.extract_text())
            
            return "".join(text_parts), metadata
        except ImportError:
            raise Exception("PyPDF2 required: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"Failed to extract PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX with metadata"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "title": "",
                "author": (
                    doc.core_properties.author
                    if hasattr(doc, "core_properties")
                    else ""
                ),
            }

            for para in doc.paragraphs:
                text += para.text + "\n"

            return text, metadata
        except ImportError:
            raise Exception("python-docx required: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_path: str) -> Tuple[str, Dict]:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            metadata = {
                "lines": len(text.split("\n")),
                "characters": len(text),
                "words": len(text.split()),
            }

            return text, metadata
        except Exception as e:
            raise Exception(f"Failed to read TXT: {str(e)}")

    @staticmethod
    def extract_csv_data(file_path: str) -> Tuple[str, Dict]:
        """Extract and summarize CSV data"""
        try:
            import csv

            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                rows = list(reader)

            metadata = {
                "rows": len(rows),
                "columns": len(rows[0]) if rows else 0,
                "headers": rows[0] if rows else [],
            }

            # Create text representation
            text = f"CSV File with {len(rows)} rows and {len(rows[0]) if rows else 0} columns\n"
            text += f"Headers: {', '.join(rows[0])}\n\n"
            text += "First 10 rows:\n"
            for row in rows[1:11]:
                text += ", ".join(row) + "\n"

            return text, metadata
        except Exception as e:
            raise Exception(f"Failed to read CSV: {str(e)}")


class AudioProcessor:
    """Advanced audio processing"""

    @staticmethod
    def get_audio_duration(file_path: str) -> float:
        """Get audio duration in seconds"""
        try:
            import librosa

            duration, _ = librosa.get_samplerate(file_path)
            return duration
        except:
            try:
                from pydub import AudioSegment

                audio = AudioSegment.from_file(file_path)
                return len(audio) / 1000.0  # Convert to seconds
            except Exception as e:
                raise Exception(f"Failed to get audio duration: {str(e)}")

    @staticmethod
    def extract_audio_features(file_path: str) -> Dict[str, Any]:
        """Extract audio features"""
        try:
            import librosa
            import numpy as np

            y, sr = librosa.load(file_path)

            features = {
                "sample_rate": sr,
                "duration": librosa.get_duration(y=y, sr=sr),
                "rms_energy": float(np.mean(librosa.feature.rms(y=y))),
                "zero_crossing_rate": float(
                    np.mean(librosa.feature.zero_crossing_rate(y)[0])
                ),
            }

            # MFCC (Mel-frequency cepstral coefficients)
            mfcc = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
            features["mfcc_mean"] = float(np.mean(mfcc))

            return features
        except ImportError:
            return {"warning": "librosa not installed for advanced audio analysis"}
        except Exception as e:
            return {"error": str(e)}

    @staticmethod
    def convert_audio_format(input_path: str, output_format: str = "wav") -> bytes:
        """Convert audio to different format"""
        try:
            from pydub import AudioSegment

            audio = AudioSegment.from_file(input_path)
            buffer = io.BytesIO()
            audio.export(buffer, format=output_format)
            return buffer.getvalue()
        except Exception as e:
            raise Exception(f"Failed to convert audio: {str(e)}")

    @staticmethod
    def normalize_audio(file_path: str) -> bytes:
        """Normalize audio level"""
        try:
            from pydub import AudioSegment

            audio = AudioSegment.from_file(file_path)

            # Normalize to -20dBFS
            target_dBFS = -20.0
            loudness = audio.dBFS
            normalizing_gain = target_dBFS - loudness
            normalized = audio.apply_gain(normalizing_gain)

            buffer = io.BytesIO()
            format_type = audio.format if hasattr(audio, "format") else "wav"
            normalized.export(buffer, format=format_type)
            return buffer.getvalue()
        except Exception as e:
            raise Exception(f"Failed to normalize audio: {str(e)}")


class VideoProcessor:
    """Advanced video processing"""

    @staticmethod
    def get_video_info(file_path: str) -> Dict[str, Any]:
        """Extract video information"""
        try:
            from moviepy.editor import VideoFileClip

            video = VideoFileClip(file_path)

            info = {
                "duration": video.duration,
                "fps": video.fps,
                "resolution": video.size,
                "width": video.w,
                "height": video.h,
                "frame_count": int(video.fps * video.duration),
            }

            video.close()
            return info
        except Exception as e:
            return {"error": str(e)}

    @staticmethod
    def extract_frames(file_path: str, num_frames: int = 5) -> List[bytes]:
        """Extract key frames from video"""
        try:
            import numpy as np
            from moviepy.editor import VideoFileClip
            from PIL import Image

            video = VideoFileClip(file_path)
            duration = video.duration

            frames_bytes = []
            for i in range(num_frames):
                t = (i * duration) / num_frames
                frame = video.get_frame(t)

                # Convert numpy array to PIL Image
                img = Image.fromarray(np.uint8(frame))

                # Convert to bytes
                buffer = io.BytesIO()
                img.save(buffer, format="PNG")
                frames_bytes.append(buffer.getvalue())

            video.close()
            return frames_bytes
        except Exception as e:
            raise Exception(f"Failed to extract frames: {str(e)}")

    @staticmethod
    def create_video_thumbnail(file_path: str, time_offset: float = 5.0) -> bytes:
        """Create thumbnail from video at specific time"""
        try:
            import numpy as np
            from moviepy.editor import VideoFileClip
            from PIL import Image

            video = VideoFileClip(file_path)

            # Ensure time_offset is within video duration
            time_offset = min(time_offset, video.duration - 0.1)

            frame = video.get_frame(time_offset)
            img = Image.fromarray(np.uint8(frame))

            # Resize to thumbnail size
            img.thumbnail((320, 240), Image.Resampling.LANCZOS)

            buffer = io.BytesIO()
            img.save(buffer, format="PNG")

            video.close()
            return buffer.getvalue()
        except Exception as e:
            raise Exception(f"Failed to create thumbnail: {str(e)}")

    @staticmethod
    def compress_video(file_path: str, quality: str = "medium") -> bytes:
        """Compress video for faster processing"""
        try:
            from moviepy.editor import VideoFileClip
<<<<<<< HEAD

            quality_settings = {
                "low": {"bitrate": "500k", "fps": 15},
                "medium": {"bitrate": "2000k", "fps": 24},
                "high": {"bitrate": "5000k", "fps": 30},
            }

            settings = quality_settings.get(quality, quality_settings["medium"])

=======
            
>>>>>>> origin/code-quality-refactor-17423438479402428749
            video = VideoFileClip(file_path)

            # This is a placeholder - actual compression would require ffmpeg integration
            # For now, just return the original file info

            video.close()

            with open(file_path, "rb") as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Failed to compress video: {str(e)}")


class MultimodalManager:
    """Central manager for all multimodal processing"""

    def __init__(self):
        self.processed_files: List[MediaFile] = []
        self.processing_history: List[Dict] = []

    def process_file(
        self, file_path: str, file_type: Optional[str] = None
    ) -> MediaFile:
        """Process uploaded file based on type"""
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        # Determine file type
        ext = os.path.splitext(filename)[1].lower().lstrip(".")

        if file_type is None:
            # Auto-detect
            if ext in [fmt.value for fmt in ImageFormat]:
                file_type = MediaType.IMAGE
            elif ext in [fmt.value for fmt in DocumentFormat]:
                file_type = MediaType.DOCUMENT
            elif ext in [fmt.value for fmt in AudioFormat]:
                file_type = MediaType.AUDIO
            elif ext in [fmt.value for fmt in VideoFormat]:
                file_type = MediaType.VIDEO
            else:
                file_type = MediaType.TEXT

        # Process based on type
        metadata = {}

        if file_type == MediaType.IMAGE:
            metadata = ImageProcessor.extract_image_metadata(file_path)
        elif file_type == MediaType.DOCUMENT:
            if ext == "pdf":
                text, doc_meta = DocumentProcessor.extract_text_from_pdf(file_path)
                metadata = doc_meta
            elif ext == "docx":
                text, doc_meta = DocumentProcessor.extract_text_from_docx(file_path)
                metadata = doc_meta
            elif ext == "csv":
                text, doc_meta = DocumentProcessor.extract_csv_data(file_path)
                metadata = doc_meta
            else:
                text, doc_meta = DocumentProcessor.extract_text_from_txt(file_path)
                metadata = doc_meta
        elif file_type == MediaType.AUDIO:
            metadata = AudioProcessor.extract_audio_features(file_path)
        elif file_type == MediaType.VIDEO:
            metadata = VideoProcessor.get_video_info(file_path)

        media_file = MediaFile(
            filename=filename,
            media_type=file_type,
            format_type=ext,
            file_size=file_size,
            processed_at=datetime.now().isoformat(),
            metadata=metadata,
        )

        self.processed_files.append(media_file)

        # Log processing
        self.processing_history.append(
            {
                "filename": filename,
                "type": file_type.value,
                "timestamp": datetime.now().isoformat(),
            }
        )

        return media_file

    def get_processing_summary(self) -> Dict[str, Any]:
        """Get summary of processed files"""
        return {
            "total_files": len(self.processed_files),
            "files_by_type": {
                "images": len(
                    [f for f in self.processed_files if f.media_type == MediaType.IMAGE]
                ),
                "documents": len(
                    [
                        f
                        for f in self.processed_files
                        if f.media_type == MediaType.DOCUMENT
                    ]
                ),
                "audio": len(
                    [f for f in self.processed_files if f.media_type == MediaType.AUDIO]
                ),
                "video": len(
                    [f for f in self.processed_files if f.media_type == MediaType.VIDEO]
                ),
            },
            "total_size": sum(f.file_size for f in self.processed_files),
            "recent_files": [f.to_dict() for f in self.processed_files[-5:]],
        }

    def clear_history(self):
        """Clear processing history"""
        self.processed_files = []
        self.processing_history = []


# Export utilities
__all__ = [
    "MediaType",
    "ImageFormat",
    "DocumentFormat",
    "AudioFormat",
    "VideoFormat",
    "MediaFile",
    "ImageProcessor",
    "DocumentProcessor",
    "AudioProcessor",
    "VideoProcessor",
    "MultimodalManager",
]
